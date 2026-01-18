
from docx import Document
import os

def inspect_awl():
    path = os.path.join("wordlist", "Headwords-of-the-Academic-Word-List.docx")
    try:
        doc = Document(path)
        print(f"Total Paragraphs: {len(doc.paragraphs)}")
        for i, para in enumerate(doc.paragraphs):
            txt = para.text.strip()
            if txt:
                print(f"{i}: {txt}")
            if i > 50: break # Inspect first 50 lines
    except Exception as e:
        print(e)

if __name__ == "__main__":
    inspect_awl()
