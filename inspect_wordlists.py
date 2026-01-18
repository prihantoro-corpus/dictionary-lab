
import os
import csv
from pypdf import PdfReader
from docx import Document

BASE_DIR = "wordlist"

def peek_csv(filename):
    print(f"\n--- {filename} ---")
    try:
        with open(os.path.join(BASE_DIR, filename), 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            rows = []
            for i, row in enumerate(reader):
                if i < 5: rows.append(row)
                else: break
            for r in rows: print(r)
    except Exception as e:
        print(f"Error: {e}")

def peek_docx(filename):
    print(f"\n--- {filename} ---")
    try:
        doc = Document(os.path.join(BASE_DIR, filename))
        # Print first 10 paragraphs
        count = 0
        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt:
                print(f"P: {txt[:100]}...")
                count += 1
            if count >= 10: break
        
        # Check tables too
        print("Tables snippet:")
        if doc.tables:
            for row in doc.tables[0].rows[:3]:
                print([cell.text.strip() for cell in row.cells])
    except Exception as e:
        print(f"Error: {e}")

def peek_pdf(filename):
    print(f"\n--- {filename} ---")
    try:
        reader = PdfReader(os.path.join(BASE_DIR, filename))
        print(f"Pages: {len(reader.pages)}")
        # Text from page 1
        if len(reader.pages) > 0:
            print(reader.pages[0].extract_text()[:500])
    except Exception as e:
        print(f"Error: {e}")

files = [
    "NGSL_1.2_stats.csv",
    "Headwords-of-the-Academic-Word-List.docx",
    "vocabulary-and-the-CEFR.docx",
    "general-service-list-headwords.pdf"
]

for f in files:
    if f.endswith('.csv'): peek_csv(f)
    elif f.endswith('.docx'): peek_docx(f)
    elif f.endswith('.pdf'): peek_pdf(f)
